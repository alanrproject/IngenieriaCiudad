from docx import Document
from PIL import Image
import io

class WordReader:
    def __init__(self, filepath):
        self.filepath = filepath

    def read(self):
        doc = Document(self.filepath)

        # Read text elements
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Read headings
        headings = [heading.text for heading in doc.paragraphs if heading.style.name.startswith('Heading')]

        # Read images
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_data = rel.target_part.blob
                img = Image.open(io.BytesIO(img_data))
                images.append(img)

        return text, headings, images


