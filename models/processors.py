import pandas as pd
from wordreader import WordReader
from memoryreader import MemoryReader
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocumentProcessor:
    def __init__(self, word_path, sheetnames, image_dict):
        self.word_reader = WordReader(word_path)
        self.memory_reader = MemoryReader(sheetnames)
        self.sheetnames = sheetnames
        self.image_dict = image_dict

    def process_document(self):
        # Leer documento word base
        doc = self.word_reader.read()  # Utiliza un método para leer el documento como objeto Document de python-docx        
        # Reemplazar valores en los párrafos de texto
        self.replace_text(doc)
        # Reemplazar valores en las tablas
        self.replace_in_tables(doc)    
        # Reemplazar imágenes en el documento
        self.replace_images(doc)



        # Guardar el documento actualizado
        output_docx = 'processed/PE1126_updated.docx'
        doc.save(output_docx)
        print(f"Archivo '{output_docx}' generado con éxito.")

    def replace_text(self, doc):
        paragraphs = doc.paragraphs
        df = pd.DataFrame(self.memory_reader.read())

        # Reemplazar valores en los párrafos de texto
        for _, row in df.iterrows():
            variable_name = row['Nombre de la variable']
            value = row['Valor']
            for paragraph in paragraphs:
                if variable_name in paragraph.text:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)  # Set the font size to 12 point
                        run.text = run.text.replace(variable_name, str(value))

    def replace_in_tables(self, doc):
        for sheetname, placeholder in self.sheetnames.items():
            table = self.memory_reader.get_tables(sheetname)
            
            if table is not None:
                # Search and replace placeholder with table
                for paragraph in doc.paragraphs:
                    if placeholder in paragraph.text:
                        # Clear existing content in the paragraph runs
                        for run in paragraph.runs:
                            run.text = ''
                        
                        # Insert table after the paragraph
                        paragraph._element.getparent().insert(paragraph._element.getparent().index(paragraph._element) + 1, table)

    def replace_images(self, doc):
        for placeholder, image_path in self.image_dict.items():
            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    # Create a new paragraph for the image
                    new_paragraph = paragraph.insert_paragraph_before()
                    new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align the paragraph
                    run = new_paragraph.add_run()
                    run.add_picture(image_path, width=Inches(4))  # Adjust the width as needed
